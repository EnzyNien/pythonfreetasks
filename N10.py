import re
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from openpyxl import Workbook


class TrasformDocxToExcel():
	#test_regex = r'([\w\d\b]+[^)])|^([a-z]{1}([)])(.+))'
	regex_dict_ = {'n':'[0-9]','l':'[a-z]','L':'[A-Z]'}

	def makeExcel(self, excel_name = ''):
		print(f'Start writing in {excel_name} ===>')
		wb = Workbook()
		sheet = wb.active

		for i,q in enumerate(self._data):
			sheet.cell(row=i+1, column=2).value = q.get('question',"")
			for j,answer in enumerate(q.get('answers',[])):
				text = '*' + answer[1] if answer[2] else answer[1]
				sheet.cell(row=i+1, column=j+6).value = text

		wb.save(excel_name)
		print(f'<=== Writing in {excel_name} is finished')
		return self

	def parsDocx(self, resp_sep = ')', resp_type = 'l'):
		''' 
		resp_sep = '.' =>	1. ...some text 2. ...some text
		resp_sep = ')' =>	1) ...some text 2) ...some text

		resp_type = n =>	1. ...some text 2. ...some text
		resp_type = l =>	a. ...some text b. ...some text
		resp_type = L =>	A. ...some text B. ...some text

		resp_color - must be HEX
		'''
		self._data = []
		doc = Document(self.doc_name)
		if not doc.paragraphs:
			err = 'File not found or wrong file name'
			raise ValueError(err)
			return self
		
		print(f'Open {self.doc_name}')

		_type = TrasformDocxToExcel.regex_dict_.get(resp_type,'[0-9]')
		_A	= '{1}'
		regex_text = rf'^({_type}+)[{resp_sep}]{_A}\s*(.*)'#.format(_type,resp_sep)
		regex = re.compile(regex_text)
		
		question_text = ''
		answers_arr = []
		print(f'Start parsing {self.doc_name} ===>')
		for par in doc.paragraphs:
			result = regex.match(par.text)
			if par.text.strip() == '':
				if answers_arr:
					self._data.append({'question':question_text,'answers':answers_arr})
				answers_arr = []
				question_text = ''
				continue

			if result is None:
				question_text += ' ' + par.text
			else:
				count, resp = result.groups()
				correct = True if (par.runs[0].font.highlight_color) == WD_COLOR_INDEX.YELLOW else False
				answers_arr.append([count,resp,correct])	
		print(f'<=== Parding {self.doc_name} is finished')
		return self

	def __init__(self, doc_name = None):
		self.doc_name = doc_name

if __name__== '__main__':
	Cls = TrasformDocxToExcel('Test.docx')
	Cls.parsDocx(resp_type='l').makeExcel('result.xlsx')
